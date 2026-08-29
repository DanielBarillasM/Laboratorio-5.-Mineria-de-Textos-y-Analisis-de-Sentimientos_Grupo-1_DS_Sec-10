"""Genera una ficha DOCX estética para presentar el repositorio del laboratorio."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "ficha_repositorio" / "Ficha_Repositorio_Laboratorio_5.docx"
REPO_URL = "https://github.com/DanielBarillasM/Laboratorio-5.-Mineria-de-Textos-y-Analisis-de-Sentimientos_Grupo-1_DS_Sec-10"
NAVY, BLUE, TEAL, MUTED, WHITE = "102A43", "2563EB", "0F9D91", "52667A", "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    relationship = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color_node, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    link.append(run)
    paragraph._p.append(link)


def add_heading(document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)


def set_cell_text(cell, text: str, *, bold=False, color=NAVY, size=10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def main() -> int:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(NAVY)

    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(7.1)
    cell = banner.cell(0, 0)
    shade(cell, NAVY)
    margins(cell, 330, 280, 330, 280)
    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LABORATORIO 5")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Minería de Textos y Análisis de Sentimientos")
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("DCEBFF")
    status = cell.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status.add_run("ENTREGA FINAL · GRUPO 1")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("6EE7B7")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(8)
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("Universidad del Valle de Guatemala  ·  Data Science  ·  Sección 10  ·  2026")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_heading(document, "Repositorio oficial")
    repo_box = document.add_table(rows=1, cols=1)
    repo_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    repo_cell = repo_box.cell(0, 0)
    shade(repo_cell, "EFF6FF")
    margins(repo_cell, 230, 240, 230, 240)
    paragraph = repo_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Código, notebook, informe y resultados reproducibles\n")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    hyperlink(paragraph, REPO_URL, REPO_URL)

    add_heading(document, "Descripción")
    paragraph = document.add_paragraph(
        "Proyecto de procesamiento de lenguaje natural que analiza 7,613 tweets y construye un "
        "clasificador para distinguir desastres reales de usos figurados. Incluye limpieza auditable, "
        "n-gramas, comparación de modelos, sentimiento VADER, contraste estadístico de negatividad, "
        "reentrenamiento, función para texto crudo y análisis de errores."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading(document, "Resultado destacado")
    metrics = document.add_table(rows=2, cols=4)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [("81.29%", "Exactitud"), ("0.781", "F1 desastre"), ("0.809", "F1 macro"), ("0.867", "ROC–AUC")]
    for col, (value, label) in enumerate(values):
        shade(metrics.cell(0, col), BLUE)
        set_cell_text(metrics.cell(0, col), value, bold=True, color=WHITE, size=15)
        metrics.cell(0, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(metrics.cell(1, col), "F8FAFC")
        set_cell_text(metrics.cell(1, col), label, bold=True, color=MUTED, size=9)
        metrics.cell(1, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    conclusion = document.add_paragraph()
    conclusion.paragraph_format.space_before = Pt(8)
    run = conclusion.add_run("Hallazgo final: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    conclusion.add_run(
        "los tweets de desastre son más negativos, pero incorporar esa variable reduce el F1. "
        "El modelo definitivo conserva TF–IDF con regresión logística."
    )

    add_heading(document, "Integrantes")
    members = document.add_table(rows=3, cols=2)
    members.alignment = WD_TABLE_ALIGNMENT.CENTER
    people = [
        ("Jorge Gabriel Palacios Sales", "231385"),
        ("Pablo Daniel Barillas Moreno", "22193"),
        ("Roberto Emiliano Otoniel", "23968"),
    ]
    for row, (name, student_id) in enumerate(people):
        shade(members.cell(row, 0), "F8FAFC" if row % 2 == 0 else "FFFFFF")
        shade(members.cell(row, 1), "F8FAFC" if row % 2 == 0 else "FFFFFF")
        set_cell_text(members.cell(row, 0), name, bold=True)
        set_cell_text(members.cell(row, 1), student_id, color=MUTED)
        members.cell(row, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_heading(document, "Contenido del repositorio")
    content = document.add_table(rows=3, cols=2)
    content.alignment = WD_TABLE_ALIGNMENT.CENTER
    entries = [
        ("Análisis", "Notebook final ejecutado, tablas y figuras"),
        ("Entrega", "Informe final PDF/LaTeX y ficha del repositorio"),
        ("Reproducibilidad", "Código modular, modelo persistido, metadatos y pruebas"),
    ]
    for row, (label, value) in enumerate(entries):
        shade(content.cell(row, 0), TEAL)
        set_cell_text(content.cell(row, 0), label, bold=True, color=WHITE)
        set_cell_text(content.cell(row, 1), value)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(12)
    run = footer.add_run("Laboratorio 5 · Grupo 1 · Data Science, Sección 10")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Ficha creada: {OUTPUT.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
